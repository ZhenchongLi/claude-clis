from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
    from docx.document import Document as DocxDocument
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordReaderError(Exception):
    pass


class WordReader:
    def __init__(self) -> None:
        if not DOCX_AVAILABLE:
            raise WordReaderError(
                "python-docx is not available. Install with: pip install python-docx"
            )

    def read_docx(self, file_path: Path | str) -> str:
        """Read DOCX content and convert to structured text"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise WordReaderError(f"File not found: {file_path}")
        
        if not file_path.suffix.lower() in ['.docx', '.doc']:
            raise WordReaderError(f"Not a Word document: {file_path}")
        
        try:
            doc = Document(str(file_path))
            content = self._extract_content(doc)
            return content
        except Exception as e:
            raise WordReaderError(f"Failed to read Word document: {str(e)}") from e

    def _extract_content(self, doc: DocxDocument) -> str:
        """Extract content from Word document preserving structure"""
        content_parts = []
        
        for element in doc.element.body:
            if isinstance(element, CT_P):
                paragraph = Paragraph(element, doc)
                text = self._process_paragraph(paragraph)
                if text:
                    content_parts.append(text)
            elif isinstance(element, CT_Tbl):
                table = Table(element, doc)
                table_md = self._process_table(table)
                if table_md:
                    content_parts.append(table_md)
        
        return "\n\n".join(content_parts)

    def _process_paragraph(self, paragraph: Paragraph) -> str:
        """Process a paragraph and return formatted text"""
        text = paragraph.text.strip()
        if not text:
            return ""
        
        # Detect heading levels based on style
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        
        if "heading 1" in style_name or "title" in style_name:
            return f"# {text}"
        elif "heading 2" in style_name:
            return f"## {text}"
        elif "heading 3" in style_name:
            return f"### {text}"
        elif "heading 4" in style_name:
            return f"#### {text}"
        elif "heading 5" in style_name:
            return f"##### {text}"
        elif "heading 6" in style_name:
            return f"###### {text}"
        else:
            # Check for bold/italic formatting
            formatted_text = self._format_runs(paragraph)
            return formatted_text

    def _format_runs(self, paragraph: Paragraph) -> str:
        """Process paragraph runs to preserve formatting"""
        formatted_parts = []
        
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # Apply formatting
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            
            formatted_parts.append(text)
        
        return "".join(formatted_parts)

    def _process_table(self, table: Table) -> str:
        """Convert Word table to Markdown table"""
        if not table.rows:
            return ""
        
        # Extract table data
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace("\n", " ")
                row_data.append(cell_text)
            table_data.append(row_data)
        
        if not table_data:
            return ""
        
        # Generate Markdown table
        markdown_lines = []
        
        # Header row
        header = table_data[0]
        markdown_lines.append("| " + " | ".join(header) + " |")
        markdown_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        
        # Data rows
        for row in table_data[1:]:
            # Ensure row has same number of columns as header
            while len(row) < len(header):
                row.append("")
            markdown_lines.append("| " + " | ".join(row[:len(header)]) + " |")
        
        return "\n".join(markdown_lines)

    def get_document_info(self, file_path: Path | str) -> dict[str, str | int]:
        """Get document metadata and information"""
        file_path = Path(file_path)
        
        try:
            doc = Document(str(file_path))
            core_props = doc.core_properties
            
            # Count elements
            paragraphs = len(doc.paragraphs)
            tables = len(doc.tables)
            
            info = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "paragraphs": paragraphs,
                "tables": tables,
            }
            
            return info
            
        except Exception as e:
            raise WordReaderError(f"Failed to get document info: {str(e)}") from e

    def is_docx_readable(self, file_path: Path | str) -> bool:
        """Check if DOCX can be read"""
        try:
            doc = Document(str(file_path))
            # Try to access basic properties
            return len(doc.paragraphs) >= 0
        except Exception:
            return False